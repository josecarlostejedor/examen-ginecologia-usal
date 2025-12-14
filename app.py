import streamlit as st
import openai
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import io
import random
from PIL import Image

# --- CONFIGURACIÓN DE LA PÁGINA ---
st.set_page_config(page_title="Generador Exámenes USAL - Pro", layout="wide")

st.markdown("""
    <style>
    .stTextArea textarea { font-size: 16px !important; font-family: 'Arial'; }
    .status-ok { color: green; font-weight: bold; }
    .justification-box {
        background-color: #f0f2f6;
        padding: 15px;
        border-radius: 10px;
        border-left: 5px solid #ff4b4b;
        margin-top: 10px;
    }
    .metric-card {
        background-color: #e6f3ff;
        padding: 10px;
        border-radius: 5px;
        text-align: center;
    }
    </style>
""", unsafe_allow_html=True)

# --- INICIALIZACIÓN DE ESTADO (SESSION STATE) ---
# Esto es vital para que no fallen las pestañas
if 'files_data' not in st.session_state:
    st.session_state['files_data'] = {} 
if 'files_processed_names' not in st.session_state:
    st.session_state['files_processed_names'] = []
if 'questions_db' not in st.session_state:
    st.session_state['questions_db'] = {}
if 'final_exam_questions' not in st.session_state:
    st.session_state['final_exam_questions'] = []

# --- FUNCIONES DE LÓGICA ---

def extract_content_robust(file):
    """Extrae Texto e IMÁGENES de PDF/PPT"""
    try:
        reader = PdfReader(file)
        text = ""
        extracted_images = []
        
        for page in reader.pages:
            t = page.extract_text()
            if t: text += t + "\n"
            try:
                for img_file_obj in page.images:
                    extracted_images.append(img_file_obj.data)
            except: pass
        
        if len(text.strip()) < 50:
            return None, [], "⚠️ PDF sin texto reconocible"
            
        return text, extracted_images, "OK"
    except Exception as e:
        return None, [], f"❌ Error: {str(e)}"

def call_openai_generator(api_key, text, na, nb, nc, topic):
    client = openai.OpenAI(api_key=api_key)
    
    # --- PROMPT ESTILO MIR (Narrativa fluida) ---
    system_prompt = """
    Eres un Catedrático de Obstetricia y Ginecología experto en redacción de preguntas tipo MIR.
    
    OBJETIVO:
    Generar preguntas de alta calidad técnica, discriminatorias y ajustadas a la realidad clínica.
    
    INSTRUCCIONES:
    1. TIPO A (Conocimiento Directo): Definiciones, anatomía, clasificaciones.
    2. TIPO B (Conocimiento Integrado): Fisiopatología, farmacología.
    3. TIPO C (CASOS CLÍNICOS - ESTILO MIR):
       - FORMATO: Redacta un ÚNICO PÁRRAFO narrativo y cohesivo. NO uses listas ni apartados.
       - CONTENIDO: Integra perfil (edad, paridad), motivo consulta, exploración y pruebas.
       - SELECCIÓN: Incluye SOLO datos relevantes (positivos y negativos) para el diagnóstico.
       - REALISMO: Usa valores numéricos (ej: "Hb 9.2 g/dL", "TA 85/50 mmHg").
       - IMÁGENES: Si aplica, escribe asumiendo que el alumno ve la imagen (ej: "...se observa la siguiente imagen:").
    
    FORMATO JSON:
    {
        "questions": [
            {
                "type": "Tipo A/B/C",
                "question": "Texto de la pregunta...",
                "options": ["a) ...", "b) ...", "c) ...", "d) ..."],
                "answer_index": 0,
                "justification": "Explicación detallada..."
            }
        ]
    }
    """
    
    user_prompt = f"Tema: {topic}. Genera: {na} Tipo A, {nb} Tipo B, {nc} Tipo C (Estilo MIR).\nTEXTO BASE:\n{text[:25000]}..."

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
            response_format={"type": "json_object"}, temperature=0.7
        )
        return json.loads(response.choices[0].message.content).get("questions", [])
    except Exception as e:
        st.error(f"Error OpenAI: {e}")
        return []

def create_header(doc, is_exam=False):
    # Cabecera Institucional
    table = doc.add_table(1, 2)
    table.autofit = False
    table.columns[0].width = Inches(4)
    table.columns[1].width = Inches(2.5)
    
    c1 = table.cell(0, 0).paragraphs[0]
    c1.add_run("VNIVERSIDAD\nD SALAMANCA\n").bold = True
    c1.runs[0].font.size = Pt(14)
    c1.add_run("CAMPUS DE EXCELENCIA INTERNACIONAL").font.size = Pt(7)
    
    c2 = table.cell(0, 1).paragraphs[0]
    c2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c2.add_run("FACULTAD DE MEDICINA\nDEPARTAMENTO DE OBSTETRICIA Y GINECOLOGÍA").bold = True
    c2.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    if is_exam:
        # Datos Alumno
        p = doc.add_paragraph()
        p.add_run("CURSO _3º____\n").bold = True
        p.add_run("APELLIDOS _________________________________________________________________________\n")
        p.add_run("NOMBRE __________________________________________ DNI _______________________")
        doc.add_paragraph("")

        # Título
        tit = doc.add_heading("Ginecología", 0)
        tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Cuadro de Instrucciones (TEXTO COMPLETO SOLICITADO)
        t = doc.add_table(1, 1)
        t.style = 'Table Grid'
        
        full_text = (
            "Lea atentamente cada cuestión antes de responder.\n"
            "Dispone de 50 minutos para responder a 40 preguntas tipo test "
            "con 4 opciones, de las que sólo una es verdadera.\n"
            "Cada pregunta correcta suma 1 punto. Las respuestas incorrectas "
            "restan 0.25 puntos. Las preguntas no contestadas no suman ni "
            "restan puntuación.\n"
            "Para aprobar el examen será necesario obtener como mínimo una "
            "puntuación final de 5 puntos.\n"
            "La valoración final en las calificaciones será sobre 10 puntos."
        )
        
        cell = t.cell(0,0)
        p_instr = cell.paragraphs[0]
        run_instr = p_instr.add_run(full_text)
        run_instr.font.size = Pt(10)
        doc.add_paragraph("")

def add_image_to_doc(doc, q):
    if 'image_data' in q and q['image_data'] is not None:
        try:
            image_stream = io.BytesIO(q['image_data'])
            doc.add_picture(image_stream, width=Inches(3.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except: pass

def create_exam_docx(questions):
    doc = Document()
    create_header(doc, is_exam=True)
    
    for i, q in enumerate(questions):
        p = doc.add_paragraph()
        p.add_run(f"{i+1}. {q['question']}").bold = True
        add_image_to_doc(doc, q)
        
        letters = ["a)", "b)", "c)", "d)"]
        for j, opt in enumerate(q['options']):
            clean = opt.split(') ', 1)[-1] if ')' in opt[:4] else opt
            doc.add_paragraph(f"{letters[j]} {clean}")
        doc.add_paragraph()
    return doc

def create_solution_docx(questions):
    doc = Document()
    create_header(doc, is_exam=False)
    doc.add_heading("SOLUCIONARIO DEL EXAMEN", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Este documento es confidencial para uso del profesorado.\n")
    
    for i, q in enumerate(questions):
        # Enunciado
        p = doc.add_paragraph()
        p.add_run(f"PREGUNTA {i+1}: ").bold = True
        p.add_run(q['question'])
        add_image_to_doc(doc, q)
        
        # Opciones
        letters = ["a)", "b)", "c)", "d)"]
        for j, opt in enumerate(q['options']):
            clean = opt.split(') ', 1)[-1] if ')' in opt[:4] else opt
            p_opt = doc.add_paragraph()
            run_opt = p_opt.add_run(f"{letters[j]} {clean}")
            if j == q['answer_index']:
                run_opt.bold = True
                run_opt.font.color.rgb = RGBColor(0, 128, 0) # Verde
                p_opt.add_run("  [CORRECTA]").bold = True
        
        # Justificación
        p_just = doc.add_paragraph()
        p_just.paragraph_format.left_indent = Inches(0.5)
        run_j = p_just.add_run(f"Justificación: {q.get('justification', 'Sin justificación.')}")
        run_j.italic = True
        run_j.font.color.rgb = RGBColor(100, 100, 100) # Gris
        doc.add_paragraph("-" * 30)
    return doc

# --- SIDEBAR ---
with st.sidebar:
    st.image("https://upload.wikimedia.org/wikipedia/commons/thumb/6/62/Escudo_de_la_Universidad_de_Salamanca.svg/1200px-Escudo_de_la_Universidad_de_Salamanca.svg.png", width=80)
    st.title("Generador USAL")
    api_key = st.text_input("Clave API OpenAI", type="password")
    st.info("Pestaña 1: Sube archivos.\nPestaña 2: Crea preguntas.\nPestaña 3: Examen.\nPestaña 4: Soluciones.")

# --- TABS PRINCIPALES ---
tab1, tab2, tab3, tab4 = st.tabs([
    "1️⃣ Subir Material", 
    "2️⃣ Generar y Editar", 
    "3️⃣ Componer Examen",
    "4️⃣ Solucionario"
])

# --- TAB 1: SUBIDA ---
with tab1:
    st.header("Carga de Archivos Docentes")
    uploaded = st.file_uploader("Sube PDFs o PPTs exportados", type="pdf", accept_multiple_files=True)
    if uploaded:
        new_files = [f for f in uploaded if f.name not in st.session_state['files_processed_names']]
        if new_files:
            bar = st.progress(0)
            for i, f in enumerate(new_files):
                text, imgs, status = extract_content_robust(f)
                if text:
                    st.session_state['files_data'][f.name] = {'text': text, 'images': imgs}
                    st.session_state['files_processed_names'].append(f.name)
                bar.progress((i+1)/len(new_files))
            st.rerun()
    
    n_files = len(st.session_state['files_data'])
    if n_files > 0:
        st.success(f"✅ {n_files} archivos cargados en memoria.")
    else:
        st.warning("No hay archivos cargados.")

# --- TAB 2: EDITOR ---
with tab2:
    st.header("Generación y Edición de Preguntas")
    temas = list(st.session_state['files_data'].keys())
    if not temas: 
        st.warning("⚠️ Sube archivos en la Pestaña 1 primero.")
    else:
        tema_sel = st.selectbox("Selecciona Tema:", temas)
        if tema_sel:
            st.divider()
            c1, c2, c3, c4 = st.columns([1,1,1,2])
            # Selectores aumentados a 40
            na = c1.number_input("A (Directas)", 0, 40, 2)
            nb = c2.number_input("B (Integradas)", 0, 40, 2)
            nc = c3.number_input("C (Casos)", 0, 40, 1)
            
            if c4.button("✨ Generar Preguntas", type="primary"):
                if not api_key: st.error("Falta API Key"); st.stop()
                with st.spinner("Generando preguntas Estilo MIR..."):
                    qs = call_openai_generator(api_key, st.session_state['files_data'][tema_sel]['text'], na, nb, nc, tema_sel)
                    if qs: 
                        st.session_state['questions_db'][tema_sel] = qs
                        st.success("¡Preguntas generadas! Aparecerán abajo.")
                        st.rerun() # Recargar para mostrar editor

            # EDITOR DE PREGUNTAS
            if tema_sel in st.session_state['questions_db']:
                qs = st.session_state['questions_db'][tema_sel]
                imgs_pdf = st.session_state['files_data'][tema_sel]['images']
                
                with st.form(f"form_{tema_sel}"):
                    updated_qs = []
                    for i, q in enumerate(qs):
                        st.markdown(f"**Pregunta {i+1}** ({q.get('type')})")
                        new_q = st.text_area("Enunciado", q['question'], key=f"q_{i}", height=120)
                        
                        # Gestión Imágenes
                        col_img_prev, col_img_ctrl = st.columns([1, 2])
                        current_img_data = q.get('image_data', None)
                        
                        with col_img_prev:
                            if current_img_data:
                                st.image(current_img_data, width=150, caption="Imagen Asignada")
                            else:
                                st.caption("Sin imagen")

                        with col_img_ctrl:
                            source = st.radio("Imagen:", ["Mantener", "Del PDF", "Subir", "Borrar"], key=f"src_{i}", horizontal=True, index=0)
                            final_img_data = current_img_data # Por defecto mantenemos
                            
                            if source == "Del PDF":
                                if imgs_pdf:
                                    idx = st.number_input(f"Idx (0-{len(imgs_pdf)-1})", 0, len(imgs_pdf)-1, 0, key=f"idx_{i}")
                                    final_img_data = imgs_pdf[idx]
                                    st.image(final_img_data, width=80)
                                else: st.warning("PDF sin imágenes")
                            elif source == "Subir":
                                uploaded_img = st.file_uploader("Archivo", type=['png','jpg'], key=f"upl_{i}")
                                if uploaded_img: final_img_data = uploaded_img.getvalue()
                            elif source == "Borrar":
                                final_img_data = None

                        # Opciones
                        opts = q['options']; 
                        while len(opts)<4: opts.append("")
                        c_o1, c_o2 = st.columns(2)
                        o0 = c_o1.text_input("a)", opts[0], key=f"o0_{i}"); o1 = c_o2.text_input("b)", opts[1], key=f"o1_{i}")
                        o2 = c_o1.text_input("c)", opts[2], key=f"o2_{i}"); o3 = c_o2.text_input("d)", opts[3], key=f"o3_{i}")
                        
                        c_ans, c_just = st.columns([1,3])
                        idx = c_ans.selectbox("Correcta", [0,1,2,3], index=q['answer_index'], format_func=lambda x:"abcd"[x], key=f"ans_{i}")
                        just = c_just.text_input("Justificación", q.get('justification',''), key=f"jus_{i}")
                        
                        updated_qs.append({
                            **q, 'question': new_q, 'options': [o0,o1,o2,o3], 'answer_index': idx, 
                            'justification': just, 'image_data': final_img_data
                        })
                        st.divider()
                    
                    if st.form_submit_button("💾 Guardar Cambios"):
                        st.session_state['questions_db'][tema_sel] = updated_qs
                        st.success("Cambios guardados correctamente.")

# --- TAB 3: GENERAR EXAMEN ---
with tab3:
    st.header("Componer Examen Final")
    
    # Recopilar todas las preguntas disponibles
    all_qs = []
    for qs in st.session_state['questions_db'].values():
        all_qs.extend(qs)
        
    st.markdown(f"<div class='metric-card'>Preguntas Disponibles en el Banco: <b>{len(all_qs)}</b></div>", unsafe_allow_html=True)
    
    if len(all_qs) == 0:
        st.warning("⚠️ No hay preguntas generadas. Ve a la Pestaña 2 y genera preguntas de algún tema.")
    else:
        st.write("---")
        num = st.number_input("Número de preguntas para el examen:", 1, 120, 40)
        
        if st.button("🎲 Generar Nuevo Modelo de Examen"):
            if len(all_qs) > num:
                sel = random.sample(all_qs, num)
            else:
                sel = all_qs
            random.shuffle(sel)
            st.session_state['final_exam_questions'] = sel
            st.success(f"¡Examen de {len(sel)} preguntas creado! Descárgalo abajo.")
            
        # Mostrar descarga si hay examen generado
        if st.session_state['final_exam_questions']:
            qs_exam = st.session_state['final_exam_questions']
            st.write(f"✅ **Examen Actual:** {len(qs_exam)} preguntas.")
            
            doc = create_exam_docx(qs_exam)
            bio = io.BytesIO()
            doc.save(bio)
            
            st.download_button(
                label="📄 Descargar Examen (.docx)",
                data=bio.getvalue(),
                file_name="Examen_Final_Ginecologia.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.info("Pulsa el botón 'Generar Nuevo Modelo' para crear el examen.")

# --- TAB 4: SOLUCIONARIO ---
with tab4:
    st.header("Solucionario")
    
    if not st.session_state['final_exam_questions']:
        st.warning("⚠️ Primero debes generar un examen en la Pestaña 3.")
    else:
        qs_exam = st.session_state['final_exam_questions']
        
        col_down, col_preview = st.columns([1, 3])
        
        with col_down:
            doc_sol = create_solution_docx(qs_exam)
            bio_sol = io.BytesIO()
            doc_sol.save(bio_sol)
            st.download_button(
                label="🔑 Descargar Solucionario (.docx)",
                data=bio_sol.getvalue(),
                file_name="Solucionario_Examen.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )
        
        with col_preview:
            st.subheader(f"Vista Previa ({len(qs_exam)} preguntas)")
            for i, q in enumerate(qs_exam):
                with st.expander(f"P{i+1}: {q['question'][:80]}..."):
                    st.write(f"**Enunciado:** {q['question']}")
                    if q.get('image_data'): st.image(q['image_data'], width=200)
                    st.markdown("**Opciones:**")
                    for j, opt in enumerate(q['options']):
                        if j == q['answer_index']:
                            st.markdown(f"- ✅ **{opt}**")
                        else:
                            st.markdown(f"- {opt}")
                    st.markdown(f"<div class='justification-box'><b>Justificación:</b><br>{q.get('justification', 'Sin justificación.')}</div>", unsafe_allow_html=True)
